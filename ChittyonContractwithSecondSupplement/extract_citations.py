from docx import Document
import re

def extract_citations(doc_path):
    doc = Document(doc_path)
    citations = []
    full_text = "\n".join(
        [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    )
    citation_pattern = re.compile(r"^\s*\d+\.\s+(.+)", re.MULTILINE)
    citations.extend(citation_pattern.findall(full_text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if citation_pattern.match(text):
                    citations.append(text)
    return citations